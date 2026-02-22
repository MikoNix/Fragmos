import streamlit as st
from pages.modules import util_sidebar
from docx import Document
from docxtpl import DocxTemplate
from io import BytesIO
import re
from pathlib import Path


class Template:
    def __init__(self, file):
        self.file = file
        self.name = file.name if hasattr(file, 'name') else 'output.docx'
        self.data_to_tags = {}
    
    def found_all_tags(self, document):
        """
        Возвращает список кортежей (tag_name, comment) В ПОРЯДКЕ ПОЯВЛЕНИЯ в документе.
        Поддерживает формат {{tag_name:Комментарий}} и {{tag_name}}.
        """
        tags_ordered = []
        seen = set()

        if hasattr(document, 'seek'):
            document.seek(0)

        doc = Document(document)

        tag_pattern = re.compile(r'\{\{([^}:]+?)(?::([^}]*))?\}\}')

        def extract(text):
            for match in tag_pattern.finditer(text):
                tag_name = match.group(1).strip()
                comment = match.group(2).strip() if match.group(2) else None
                if tag_name not in seen:
                    seen.add(tag_name)
                    tags_ordered.append((tag_name, comment))

        for para in doc.paragraphs:
            extract(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    extract(cell.text)

        if not tags_ordered:
            raise ValueError("Теги не найдены")

        return tags_ordered

    def get_tags_data(self, tags):
        """
        tags — список кортежей (tag_name, comment) в порядке появления
        """
        self.data_to_tags = {}
        st.markdown("### Заполните значения для тегов:")

        with st.form(key="tags_form"):
            for tag_name, comment in tags:
                label = comment if comment else tag_name
                self.data_to_tags[tag_name] = st.text_area(
                    f"**{label}**",
                    key=f"input_{tag_name}",
                    placeholder=f"Введите значение для {label}",
                    height=120,
                )

            submitted = st.form_submit_button(
                "Сгенерировать документ",
                icon=":material/auto_mode:",
                use_container_width=True
            )

        if submitted:
            empty_fields = [
                (comment if comment else tag_name)
                for tag_name, comment in tags
                if not self.data_to_tags.get(tag_name, "").strip()
            ]

            if empty_fields:
                st.warning(
                    f"Пожалуйста, заполните все поля. Пустые поля: {', '.join(empty_fields)}",
                    icon=":material/warning:"
                )
            else:
                self.render_template()

    def render_template(self):
        try:
            self.file.seek(0)

            doc = DocxTemplate(self.file)
            doc.render(self.data_to_tags)

            output = BytesIO()
            doc.save(output)
            output.seek(0)

            output_filename = f"filled_{self.name}"

            st.success("Документ успешно сгенерирован!", icon=":material/check_circle:")

            st.download_button(
                "Скачать заполненный документ",
                data=output,
                file_name=output_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                icon=":material/download:",
                use_container_width=True
            )

        except Exception as e:
            st.error(f"Ошибка при генерации документа: {str(e)}", icon=":material/error:")

    @st.dialog("Выбор шаблона", width="large")
    def show_file_picker_dialog(self):
        templates_folder = Path(__file__).parent.parent / "static" / "templates"
        templates_files = []

        if templates_folder.exists():
            templates_files = list(templates_folder.glob("*.docx"))
            templates_files = [f for f in templates_files if f.is_file()]

        uploaded_file = st.file_uploader(
            "📁 Загрузить свой файл",
            type=["docx"],
            key="dialog_file_uploader"
        )

        search_query = st.text_input(
            "Поиск шаблонов",
            placeholder="Введите название шаблона...",
            key="search_template"
        )

        if uploaded_file:
            st.session_state.selected_file = uploaded_file
            st.rerun()

        if templates_files:
            filtered_templates = templates_files
            if search_query:
                filtered_templates = [
                    f for f in templates_files
                    if search_query.lower() in f.stem.lower()
                ]

            st.markdown("### Доступные шаблоны")

            for template_file in filtered_templates:
                col1, col2 = st.columns([5, 1])
                with col1:
                    st.markdown(f"📄 **{template_file.stem}**")
                with col2:
                    if st.button("Выбрать", key=f"select_{template_file.stem}", use_container_width=True):
                        try:
                            with open(template_file, 'rb') as f:
                                file_bytes = f.read()
                                file = BytesIO(file_bytes)
                                file.name = template_file.name
                            st.session_state.selected_file = file
                            st.rerun()
                        except Exception as e:
                            st.error(f"Ошибка: {str(e)}")
        else:
            st.info("Нет доступных шаблонов", icon=":material/info:")

    def show_templates(self):
        if 'selected_file' not in st.session_state:
            st.session_state.selected_file = None

        if st.button("Выбрать или загрузить шаблон", icon=":material/folder_open:", use_container_width=True):
            self.show_file_picker_dialog()

        if st.session_state.selected_file is not None:
            st.write(f"Выбран шаблон: {st.session_state.selected_file.name}")
            if st.button("Отменить выбор шаблона", icon=":material/close:", use_container_width=True):
                st.session_state.selected_file = None
                st.rerun()
            template = Template(st.session_state.selected_file)
            template.start()

    def template_maker_ai(self):
        st.info(
            "Функция генерации шаблона с помощью ИИ будет добавлена в следующих версиях",
            icon=":material/construction:"
        )

    def start(self):
        try:
            tags = self.found_all_tags(self.file)

            with st.expander(f"Найдено тегов: {len(tags)}", expanded=False):
                st.write(", ".join([
                    f"`{tag_name}`" + (f" ({comment})" if comment else "")
                    for tag_name, comment in tags
                ]))

            self.get_tags_data(tags)

        except ValueError:
            st.warning(
                "В шаблоне не найдено тегов для замены. "
                "Пожалуйста, используйте теги в формате **{{tag_name}}** или **{{tag_name:Комментарий}}**",
                icon=":material/warning:"
            )
        except Exception as e:
            st.error(
                f"Ошибка при обработке шаблона: {str(e)}",
                icon=":material/error:"
            )
            col1, col2 = st.columns([1, 3])
            with col1:
                if st.button("Перезагрузить", icon=":material/refresh:", use_container_width=True):
                    st.rerun()


def load_css():
    css_file = Path(__file__).parent / "styles" / "modal.css"
    if css_file.exists():
        with open(css_file) as f:
            st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)


util_sidebar()
load_css()

st.markdown("## Генератор отчетов")

options = {
    0: "Сгенерировать отчет из шаблона :material/auto_mode:",
    1: "Сгенерировать шаблон с помощью ИИ :material/robot:"
}

selection = st.segmented_control(
    "Select",
    options=options.keys(),
    format_func=lambda option: options[option],
    selection_mode="single", label_visibility="collapsed",
)

templates = Template(None)

if selection == 0:
    st.info(
        "Загрузите или выберете шаблон в формате .docx, содержащий теги в формате **{{tag_name}}** или **{{tag_name:Комментарий}}**. "
        "После загрузки вы сможете ввести значения для каждого тега и сгенерировать заполненный документ.",
        icon=":material/info:"
    )
    templates.show_templates()
elif selection == 1:
    templates.template_maker_ai()