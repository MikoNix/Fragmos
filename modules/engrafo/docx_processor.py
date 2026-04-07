"""
docx_processor.py — подстановка тегов в docx-шаблон через python-docx.

Поддерживает форматы тегов:
  {{key}}          — простой тег
  {{key:Подсказка}} — тег с подсказкой (подсказка игнорируется при рендере)

Теги могут быть разбиты Word на несколько XML-runs — обрабатывается корректно.
Форматирование первого run тега сохраняется для результирующего текста.
"""

import base64
import io
import json
import os
import re
import copy
from html import unescape as _html_unescape
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


# Паттерн тега: {{key}} или {{key:подсказка}}
_TAG_RE = re.compile(r"\{\{([^}:]+)(?::[^}]*)?\}\}")

# HTML: <img src="..."> capture group 1 = src value
_IMG_SRC_RE = re.compile(r'<img[^>]+src=["\']([^"\']+)["\'][^>]*/?>',  re.IGNORECASE)
# Strip any remaining HTML tags
_HTML_STRIP_RE = re.compile(r'<[^>]+>')
# Replace <br> with newline before stripping
_BR_RE = re.compile(r'<br\s*/?>', re.IGNORECASE)
# Block-level tags that act as paragraph separators in contenteditable
_BLOCK_END_RE = re.compile(r'</(div|p|li)\s*>', re.IGNORECASE)


def _runs_to_text(runs) -> str:
    return "".join(r.text for r in runs)


def _is_image_value(value: str) -> bool:
    """Проверяет, является ли значение plain base64 data URL картинки."""
    return value.startswith("data:image/")


def _contains_inline_images(value: str) -> bool:
    """Проверяет, содержит ли HTML-значение теги <img>."""
    return "<img" in value


def _strip_html(html: str) -> str:
    """Убрать HTML-теги, заменить <br> и блочные теги на перенос строки."""
    text = _BR_RE.sub("\n", html)
    text = _BLOCK_END_RE.sub("\n", text)
    text = _HTML_STRIP_RE.sub("", text)
    # Collapse multiple consecutive newlines to max two
    text = re.sub(r'\n{3,}', '\n\n', text)
    return _html_unescape(text)


def _decode_image(data_url: str) -> tuple[io.BytesIO, str]:
    """Декодирует data URL картинки. Возвращает (BytesIO, mime_type)."""
    header, data = data_url.split(",", 1)
    mime = header.split(";")[0].split(":")[1]  # image/png
    return io.BytesIO(base64.b64decode(data)), mime


def _clear_runs(para) -> None:
    p_elem = para._p
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)


def _insert_html_with_images(para, html: str) -> None:
    """
    Разбирает HTML со встроенными <img> и вставляет в параграф:
    текстовые части — как runs, картинки — как inline picture runs.
    """
    _clear_runs(para)
    # re.split с capturing group чередует текст и src
    parts = _IMG_SRC_RE.split(html)
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Текстовая часть
            text = _strip_html(part)
            if text:
                para.add_run(text)
        else:
            # src картинки
            src = part
            if src.startswith("data:image/"):
                try:
                    img_stream, _ = _decode_image(src)
                    run = para.add_run()
                    run.add_picture(img_stream, width=Inches(4))
                except Exception:
                    pass  # пропустить картинку если не удалось декодировать


def _is_ctx_value(value: str) -> bool:
    """Проверяет, является ли значение contextualizer-сентинелом __ctx__:{...}."""
    return value.startswith("__ctx__:")


# ── Markdown → DOCX ──────────────────────────────────────────────────────────


_MD_BOLD_ITALIC_RE = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*")
_MD_LIST_BULLET_RE = re.compile(r"^[ \t]*-[ \t]+(.+)$")
_MD_LIST_NUM_RE    = re.compile(r"^[ \t]*\d+\.[ \t]+(.+)$")


def _parse_inline_markdown(text: str) -> list[tuple[str, bool, bool]]:
    """
    Разбить строку на сегменты (текст, bold, italic).
    Поддерживает ***bold+italic***, **bold**, *italic*.
    Возвращает список (chunk, is_bold, is_italic).
    """
    segments: list[tuple[str, bool, bool]] = []
    pos = 0
    for m in _MD_BOLD_ITALIC_RE.finditer(text):
        if m.start() > pos:
            segments.append((text[pos:m.start()], False, False))
        if m.group(1) is not None:
            segments.append((m.group(1), True, True))
        elif m.group(2) is not None:
            segments.append((m.group(2), True, False))
        else:
            segments.append((m.group(3), False, True))
        pos = m.end()
    if pos < len(text):
        segments.append((text[pos:], False, False))
    return segments


def _add_formatted_runs(para, text: str, base_bold: bool = False, base_italic: bool = False) -> None:
    """Добавить runs с markdown inline форматированием в параграф."""
    for chunk, bold, italic in _parse_inline_markdown(text):
        if not chunk:
            continue
        run = para.add_run(chunk)
        run.bold = base_bold or bold
        run.italic = base_italic or italic


def _add_para_with_md_after(parent_elem, ref_elem, line: str) -> object:
    """
    Создать новый параграф с markdown-форматированием после ref_elem.
    Обрабатывает bullet (-) и numbered (1.) списки.
    Возвращает новый XML-элемент параграфа.
    """
    from docx.oxml import OxmlElement

    bullet_m = _MD_LIST_BULLET_RE.match(line)
    num_m    = _MD_LIST_NUM_RE.match(line)

    p_elem = OxmlElement("w:p")
    ppr = OxmlElement("w:pPr")

    if bullet_m or num_m:
        # Отступ для списка
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ppr.append(ind)

    p_elem.append(ppr)

    # Текст элемента
    item_text = (bullet_m.group(1) if bullet_m else
                 num_m.group(1) if num_m else line)
    prefix = "• " if bullet_m else ""

    # Создаём runs вручную через OxmlElement
    for chunk, bold, italic in _parse_inline_markdown(prefix + item_text):
        if not chunk:
            continue
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        has_rpr = False
        if bold:
            b = OxmlElement("w:b")
            rpr.append(b)
            has_rpr = True
        if italic:
            i_elem = OxmlElement("w:i")
            rpr.append(i_elem)
            has_rpr = True
        if has_rpr:
            r.append(rpr)
        t = OxmlElement("w:t")
        t.text = chunk
        if chunk.startswith(" ") or chunk.endswith(" "):
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        p_elem.append(r)

    idx = list(parent_elem).index(ref_elem)
    parent_elem.insert(idx + 1, p_elem)
    return p_elem


def _insert_contextual_content(para, value: str) -> None:
    """
    Обрабатывает значение вида __ctx__:{...} от contextualizer.

    Структура payload:
        {
            "content": "текст...",
            "images": [
                {"path": "/abs/path.png", "caption": "Рисунок 1", "inline_after": "предложение."}
            ],
            "options": {"image_align": "center"}
        }

    Вставляет текст параграфами, после предложения inline_after вставляет
    изображение и подпись.
    """
    try:
        payload = json.loads(value[len("__ctx__:"):])
    except (json.JSONDecodeError, IndexError):
        _clear_runs(para)
        para.add_run(value)
        return

    content = payload.get("content", "")
    images = payload.get("images", [])
    options = payload.get("options", {})
    align = options.get("image_align", "center")

    # Определяем выравнивание
    _align_map = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    pic_alignment = _align_map.get(str(align).lower(), WD_ALIGN_PARAGRAPH.CENTER)

    # Строим карту: предложение → список изображений после него
    inline_map: dict[str, list[dict]] = {}
    for img in images:
        key = (img.get("inline_after") or "").strip()
        inline_map.setdefault(key, []).append(img)

    _clear_runs(para)
    parent_elem = para._p.getparent()
    insert_after_elem = para._p

    def _add_para_after(ref_elem, text: str = "", bold: bool = False,
                        alignment=None, font_size: int | None = None):
        """Добавить параграф после ref_elem в parent_elem."""
        from docx.oxml import OxmlElement
        p = OxmlElement("w:p")
        if alignment is not None or font_size is not None:
            ppr = OxmlElement("w:pPr")
            if alignment is not None:
                jc = OxmlElement("w:jc")
                _align_names = {
                    WD_ALIGN_PARAGRAPH.CENTER: "center",
                    WD_ALIGN_PARAGRAPH.LEFT: "left",
                    WD_ALIGN_PARAGRAPH.RIGHT: "right",
                }
                jc.set(qn("w:val"), _align_names.get(alignment, "center"))
                ppr.append(jc)
            p.append(ppr)
        if text:
            r = OxmlElement("w:r")
            if bold or font_size:
                rpr = OxmlElement("w:rPr")
                if bold:
                    b = OxmlElement("w:b")
                    rpr.append(b)
                if font_size:
                    sz = OxmlElement("w:sz")
                    sz.set(qn("w:val"), str(font_size * 2))
                    rpr.append(sz)
                r.append(rpr)
            t = OxmlElement("w:t")
            t.text = text
            if text.startswith(" ") or text.endswith(" "):
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            r.append(t)
            p.append(r)
        idx = list(parent_elem).index(ref_elem)
        parent_elem.insert(idx + 1, p)
        return p

    def _add_image_para_after(ref_elem, img_path: str, caption: str):
        """Вставить параграф с изображением и подпись после ref_elem."""
        from docx.oxml import OxmlElement
        # Изображение
        img_p = OxmlElement("w:p")
        ppr = OxmlElement("w:pPr")
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        ppr.append(jc)
        img_p.append(ppr)

        if os.path.isfile(img_path):
            try:
                img_r = OxmlElement("w:r")
                # Используем временный Document чтобы вставить картинку
                _tmp_doc = Document()
                _tmp_para = _tmp_doc.paragraphs[0]
                _run = _tmp_para.add_run()
                _run.add_picture(img_path, width=Inches(4))
                # Извлекаем drawing XML из временного документа
                drawing_elem = _run._r.find(qn("w:drawing"))
                if drawing_elem is not None:
                    img_r.append(copy.deepcopy(drawing_elem))
                    img_p.append(img_r)
            except Exception:
                r = OxmlElement("w:r")
                t = OxmlElement("w:t")
                t.text = f"[image: {os.path.basename(img_path)}]"
                r.append(t)
                img_p.append(r)

        idx = list(parent_elem).index(ref_elem)
        parent_elem.insert(idx + 1, img_p)

        # Подпись под изображением
        if caption:
            cap_p = _add_para_after(img_p, caption, alignment=pic_alignment, font_size=9)
            return cap_p
        return img_p

    # Разбиваем контент на строки: списки — отдельные элементы, остальное — по предложениям
    def _split_content_lines(text: str) -> list[str]:
        """Разбить контент на логические строки: списки и предложения."""
        result = []
        for raw_line in text.split("\n"):
            line = raw_line.strip()
            if not line:
                continue
            if _MD_LIST_BULLET_RE.match(line) or _MD_LIST_NUM_RE.match(line):
                result.append(line)
            else:
                # Обычный текст разбиваем по предложениям
                sents = re.split(r"(?<=[.!?])\s+", line)
                result.extend(s for s in sents if s.strip())
        return result

    lines = _split_content_lines(content)

    # Вставляем первую строку в исходный параграф (с markdown форматированием)
    first_line = lines[0] if lines else content
    _add_formatted_runs(para, first_line)

    # Изображения после первой строки
    after_elem = insert_after_elem
    imgs_after_first = inline_map.get(first_line.strip(), [])
    for img in imgs_after_first:
        after_elem = _add_image_para_after(after_elem, img["path"], img.get("caption", ""))

    # Остальные строки — каждая как новый параграф
    for line in lines[1:]:
        after_elem = _add_para_with_md_after(parent_elem, after_elem, line)
        imgs_after = inline_map.get(line.strip(), [])
        for img in imgs_after:
            after_elem = _add_image_para_after(after_elem, img["path"], img.get("caption", ""))

    # Изображения без привязки к конкретной строке (inline_after == "")
    for img in inline_map.get("", []):
        after_elem = _add_image_para_after(after_elem, img["path"], img.get("caption", ""))


def _replace_in_paragraph(para, tag_values: dict) -> None:
    """
    Склеивает текст всех runs параграфа, находит теги, заменяет значениями.
    Форматирование берётся от первого run тега.
    Поддерживает значения: plain text, plain data URL, HTML с <img>.
    """
    full_text = _runs_to_text(para.runs)
    if "{{" not in full_text:
        return

    # Ищем первый тег с изображением или HTML-содержимым
    for m in _TAG_RE.finditer(full_text):
        key = m.group(1).strip()
        val = tag_values.get(key, "")
        if not val:
            continue

        if _is_ctx_value(val):
            # Contextualizer-сентинел: текст + изображения + опции
            _clear_runs(para)
            _insert_contextual_content(para, val)
            return

        if _is_image_value(val):
            # Plain data URL — вставляем как картинку
            _clear_runs(para)
            run = para.add_run()
            img_stream, _ = _decode_image(val)
            try:
                run.add_picture(img_stream, width=Inches(4))
            except Exception:
                run.text = f"[image:{key}]"
            return

        if _contains_inline_images(val):
            # HTML с inline-картинками — вставляем смешанный контент
            _insert_html_with_images(para, val)
            return

    # Обычная замена текста (без картинок)
    def _resolve(m):
        v = tag_values.get(m.group(1).strip(), m.group(0))
        # Если значение содержит HTML (например <br>, <div>), strip тегов
        return _strip_html(v) if "<" in v else v

    new_text = _TAG_RE.sub(_resolve, full_text)
    if new_text == full_text:
        return

    # Сохраним форматирование первого run
    if para.runs:
        first_run = para.runs[0]
        rpr_xml = copy.deepcopy(first_run._r.find(qn("w:rPr")))
    else:
        rpr_xml = None

    # Удаляем все run-элементы из параграфа
    from docx.oxml import OxmlElement
    p_elem = para._p
    for r in p_elem.findall(qn("w:r")):
        p_elem.remove(r)

    # Разбиваем текст по строкам. Каждая строка — отдельный run,
    # между строками вставляем <w:br/> (мягкий перенос строки в Word).
    lines = new_text.split("\n")
    # Убираем пустую последнюю строку если текст заканчивается на \n
    if lines and lines[-1] == "":
        lines = lines[:-1]
    if not lines:
        lines = [""]

    for i, line in enumerate(lines):
        new_r = OxmlElement("w:r")
        if rpr_xml is not None:
            new_r.append(copy.deepcopy(rpr_xml))
        if i > 0:
            # Перед каждой строкой (кроме первой) добавляем перенос строки
            br = OxmlElement("w:br")
            new_r.append(br)
        if line:
            new_t = OxmlElement("w:t")
            new_t.text = line
            if line.startswith(" ") or line.endswith(" "):
                new_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            new_r.append(new_t)
        p_elem.append(new_r)


def render_docx(template_path: str, output_path: str, tag_values: dict) -> str:
    """
    Заполнить шаблон значениями тегов и сохранить результат.

    Args:
        template_path: путь к исходному .docx шаблону
        output_path:   путь для сохранения заполненного документа
        tag_values:    словарь {key: значение}

    Returns:
        output_path
    """
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    doc = Document(template_path)

    # Обходим все параграфы документа
    for para in doc.paragraphs:
        _replace_in_paragraph(para, tag_values)

    # Параграфы внутри таблиц
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, tag_values)

    doc.save(output_path)
    return output_path
