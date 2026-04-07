"""
file_processor.py — определение типа файла, распаковка архивов,
извлечение текста (PDF/Word) и OCR изображений.

Возвращает структурированный результат: список ProcessedFile.
Файлы НЕ сохраняются в sources/ — хранение выполняет context_builder.
"""

import io
import os
import tarfile
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import yaml

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
_CONFIG_PATH = os.path.join(_THIS_DIR, "config.yaml")

with open(_CONFIG_PATH, encoding="utf-8") as _f:
    _CFG = yaml.safe_load(_f)

_EXT = _CFG["supported_extensions"]
_OCR_LANG = _CFG["ocr"]["language"]
_MAX_TEXT = _CFG["context"]["max_text_chars"]
_MAX_DEPTH = _CFG["context"]["max_archive_depth"]


# ── Типы ──────────────────────────────────────────────────────────────────────


@dataclass
class ProcessedFile:
    """Результат обработки одного файла."""
    file_type: str           # 'pdf' | 'word' | 'image'
    original_filename: str
    raw_data: Optional[bytes] = None   # исходные байты (для изображений)
    text_content: Optional[str] = None
    ocr_text: Optional[str] = None
    warnings: list[str] = field(default_factory=list)


# ── Определение типа ──────────────────────────────────────────────────────────


def detect_type(filename: str) -> Optional[str]:
    """Вернуть 'image'|'pdf'|'word'|'archive' по расширению, или None."""
    ext = Path(filename).suffix.lower()
    for ftype, exts in _EXT.items():
        if ext in exts:
            return ftype
    return None


# ── Извлечение текста из PDF ──────────────────────────────────────────────────


def _extract_pdf_text(data: bytes) -> tuple[str, list[str]]:
    """Извлечь текст из PDF. Возвращает (text, warnings)."""
    warnings = []
    try:
        import pdfplumber
    except ImportError:
        return "", ["pdfplumber не установлен — PDF текст недоступен"]

    text_parts = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                text_parts.append(page_text)
    except Exception as e:
        warnings.append(f"Ошибка чтения PDF: {e}")
        return "", warnings

    full_text = "\n".join(text_parts).strip()

    if not full_text:
        full_text, ocr_warns = _ocr_pdf_fallback(data)
        warnings.extend(ocr_warns)

    if len(full_text) > _MAX_TEXT:
        full_text = full_text[:_MAX_TEXT]
        warnings.append(f"Текст обрезан до {_MAX_TEXT} символов")

    return full_text, warnings


def _ocr_pdf_fallback(data: bytes) -> tuple[str, list[str]]:
    """OCR fallback для сканированных PDF через pdf2image + Tesseract."""
    warnings = []
    try:
        from pdf2image import convert_from_bytes
        import pytesseract
    except ImportError:
        return "", ["pdf2image или pytesseract не установлены — OCR fallback недоступен"]

    parts = []
    try:
        images = convert_from_bytes(data, dpi=200)
        for img in images:
            text = pytesseract.image_to_string(img, lang=_OCR_LANG)
            parts.append(text)
    except Exception as e:
        warnings.append(f"OCR fallback не удался: {e}")
        return "", warnings

    return "\n".join(parts).strip(), warnings


# ── Извлечение текста из Word ─────────────────────────────────────────────────


def _extract_docx_text(data: bytes) -> tuple[str, list[str]]:
    """Извлечь текст из .docx файла."""
    warnings = []
    try:
        from docx import Document
    except ImportError:
        return "", ["python-docx не установлен"]

    try:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)
        text = "\n".join(paragraphs)
    except Exception as e:
        warnings.append(f"Ошибка чтения DOCX: {e}")
        return "", warnings

    if len(text) > _MAX_TEXT:
        text = text[:_MAX_TEXT]
        warnings.append(f"Текст обрезан до {_MAX_TEXT} символов")

    return text, warnings


# ── OCR изображения ───────────────────────────────────────────────────────────


def _ocr_image(data: bytes) -> tuple[str, list[str]]:
    """OCR изображения через Tesseract. Возвращает (text, warnings)."""
    warnings = []
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return "[OCR unavailable — pytesseract/Pillow не установлены]", []

    try:
        img = Image.open(io.BytesIO(data))
        text = pytesseract.image_to_string(img, lang=_OCR_LANG)
        return text.strip(), warnings
    except Exception as e:
        warnings.append(f"OCR не удался: {e}")
        return "[OCR unavailable]", warnings


# ── Обработка одного файла ────────────────────────────────────────────────────


def _process_single(filename: str, data: bytes) -> Optional[ProcessedFile]:
    """Обработать файл с известным типом. Вернуть ProcessedFile или None."""
    ftype = detect_type(filename)
    if ftype is None or ftype == "archive":
        return None

    warnings = []
    text_content = None
    ocr_text = None
    raw_data = None

    if ftype == "pdf":
        text_content, warnings = _extract_pdf_text(data)
    elif ftype == "word":
        text_content, warnings = _extract_docx_text(data)
    elif ftype == "image":
        ocr_text, warnings = _ocr_image(data)
        raw_data = data  # сохраняем байты для записи в папку отчёта

    return ProcessedFile(
        file_type=ftype,
        original_filename=filename,
        raw_data=raw_data,
        text_content=text_content,
        ocr_text=ocr_text,
        warnings=warnings,
    )


# ── Распаковка архива ─────────────────────────────────────────────────────────


def _safe_path(base: str, member_path: str) -> Optional[str]:
    """Tar-slip защита: возвращает None если путь выходит за base."""
    target = os.path.realpath(os.path.join(base, member_path))
    if not target.startswith(os.path.realpath(base) + os.sep):
        return None
    return target


def _process_archive(filename: str, data: bytes, depth: int) -> list[ProcessedFile]:
    """Рекурсивно распаковать архив и обработать файлы внутри."""
    if depth > _MAX_DEPTH:
        return []

    results = []
    warnings_list = []
    ext = Path(filename).suffix.lower()
    full_ext = "".join(Path(filename).suffixes[-2:]).lower()

    with tempfile.TemporaryDirectory() as tmp:
        try:
            if ext in (".7z",):
                try:
                    import py7zr
                    with py7zr.SevenZipFile(io.BytesIO(data), mode="r") as sz:
                        sz.extractall(path=tmp)
                except ImportError:
                    warnings_list.append("py7zr не установлен — 7z архивы не поддерживаются")
                    return []
                except Exception as e:
                    warnings_list.append(f"Ошибка распаковки 7z: {e}")
                    return []

            elif ext in (".tar", ".gz", ".bz2", ".tgz") or full_ext in (".tar.gz", ".tar.bz2"):
                try:
                    with tarfile.open(fileobj=io.BytesIO(data)) as tf:
                        for member in tf.getmembers():
                            if not member.isfile():
                                continue
                            safe = _safe_path(tmp, member.name)
                            if safe is None:
                                warnings_list.append(f"tar-slip пропущен: {member.name}")
                                continue
                            try:
                                tf.extract(member, tmp)
                            except Exception as e:
                                warnings_list.append(f"Не удалось извлечь {member.name}: {e}")
                except Exception as e:
                    warnings_list.append(f"Ошибка распаковки tar: {e}")
                    return []
            else:
                return []

            # Обходим извлечённые файлы
            for root, _dirs, files in os.walk(tmp):
                for fname in sorted(files):
                    fpath = os.path.join(root, fname)
                    inner_ftype = detect_type(fname)
                    if inner_ftype is None:
                        continue
                    with open(fpath, "rb") as f:
                        inner_data = f.read()
                    if inner_ftype == "archive":
                        sub = _process_archive(fname, inner_data, depth + 1)
                        results.extend(sub)
                    else:
                        pf = _process_single(fname, inner_data)
                        if pf:
                            pf.warnings.extend(warnings_list)
                            results.append(pf)

        except Exception as e:
            results.append(ProcessedFile(
                file_type="archive",
                original_filename=filename,
                warnings=[f"Ошибка распаковки архива: {e}"],
            ))

    return results


# ── Публичный интерфейс ───────────────────────────────────────────────────────


def process_upload(filename: str, data: bytes) -> tuple[list[ProcessedFile], list[str]]:
    """
    Обработать загруженный файл.
    Возвращает (список ProcessedFile, список предупреждений верхнего уровня).
    """
    ftype = detect_type(filename)
    warnings = []

    if ftype is None:
        return [], [f"Неподдерживаемый тип файла: {filename}"]

    if ftype == "archive":
        results = _process_archive(filename, data, depth=1)
        return results, warnings

    pf = _process_single(filename, data)
    if pf:
        return [pf], pf.warnings
    return [], [f"Не удалось обработать файл: {filename}"]
